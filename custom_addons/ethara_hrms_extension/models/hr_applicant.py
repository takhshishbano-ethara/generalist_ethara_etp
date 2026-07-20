import json
import logging
import re
import threading
from html import unescape
from io import BytesIO

import odoo
from odoo import SUPERUSER_ID, api, fields, models
from odoo.exceptions import UserError

_logger = logging.getLogger(__name__)

_RESUME_RECOMMENDATIONS = [
    ('shortlist', 'Shortlist'),
    ('reject', 'Reject'),
    ('maybe', 'Maybe'),
    ('needs_review', 'Needs Review'),
]
_RESUME_LLM_STATUSES = [
    ('pending', 'Pending'),
    ('processing', 'Processing'),
    ('completed', 'Completed'),
    ('failed', 'Failed'),
]


class HrApplicant(models.Model):
    _inherit = 'hr.applicant'

    candidate_id = fields.Many2one(
        'ethara.candidate', string='Candidate',
        ondelete='set null', index=True,
    )
    portfolio_url = fields.Char(string='Portfolio URL')
    github_url = fields.Char(string='GitHub URL')
    resume_url = fields.Char(string='Resume URL')
    cancellation_reason = fields.Text(string='Cancellation Reason')
    status_updated_at = fields.Datetime(string='Status Updated At')

    job_history_ids = fields.One2many(
        'hr.applicant.job.history', 'applicant_id',
        string='Job Change History',
    )

    resume_score = fields.Float(string='Resume Score', digits=(4, 1))
    resume_recommendation = fields.Selection(
        _RESUME_RECOMMENDATIONS, string='LLM Recommendation',
    )
    resume_summary = fields.Text(string='Resume Summary')
    resume_strengths = fields.Text(string='Strengths')
    resume_gaps = fields.Text(string='Gaps')
    resume_screening_payload = fields.Text(string='Screening Payload (JSON)')
    resume_llm_status = fields.Selection(
        _RESUME_LLM_STATUSES, string='LLM Status', default='pending',
    )
    resume_llm_error = fields.Text(string='LLM Error')
    resume_llm_model_used = fields.Char(string='LLM Model Used')
    resume_llm_prompt_tokens = fields.Integer(string='Prompt Tokens')
    resume_llm_completion_tokens = fields.Integer(string='Completion Tokens')
    resume_llm_latency_ms = fields.Integer(string='LLM Latency (ms)')
    resume_screened_at = fields.Datetime(string='Last Screened At')
    resume_manual_override_reason = fields.Text(string='Manual Override Reason')
    resume_manual_override_at = fields.Datetime(string='Overridden At')
    resume_manual_override_by_id = fields.Many2one(
        'res.users', string='Overridden By', ondelete='set null',
    )
    priority_score = fields.Float(string='Priority Score', digits=(4, 2))
    is_reapplication_blocked = fields.Boolean(string='Re-application Blocked')
    blacklist_reason = fields.Text(string='Blacklist Reason')
    on_hold = fields.Boolean(string='On Hold')

    pipeline_status = fields.Selection(
        [
            ('applied',     'Applied'),
            ('shortlisted', 'Shortlisted'),
            ('evaluation',  'Evaluation'),
            ('submission',  'Submission'),
            ('contract',    'Contract'),
            ('compliance',  'Compliance'),
            ('email_id',    'Email ID'),
            ('onboarded',   'Onboarded'),
            ('rejected',    'Rejected'),
        ],
        string='Pipeline Status',
        default='applied',
        index=True,
        tracking=True,
    )

    candidate_code = fields.Char(
        string='Candidate Code', index=True, copy=False,
        help='Auto-generated ETH-MMMYY-XXXXXX identifier',
    )
    source_type = fields.Selection([
        ('direct_application', 'Direct Application'),
        ('vendor', 'Vendor'),
        ('campus_hire', 'Campus Hire'),
        ('employee_referral', 'Employee Referral'),
        ('internal_hiring', 'Internal Hiring'),
        ('lateral_hiring', 'Lateral Hiring'),
    ], string='Source Type', default='direct_application')
    source_reference_id = fields.Char(string='Source Reference')
    current_company = fields.Char(string='Current Company')
    current_ctc = fields.Float(string='Current CTC (LPA)', digits=(10, 2))
    expected_ctc = fields.Float(string='Expected CTC (LPA)', digits=(10, 2))
    notice_period_days = fields.Integer(string='Notice Period (days)')
    resume_text = fields.Text(string='Extracted Resume Text')
    is_duplicate = fields.Boolean(string='Is Duplicate')
    duplicate_reason = fields.Text(string='Duplicate Reason')

    marital = fields.Selection(
        [
            ('single', 'Single'),
            ('married', 'Married'),
            ('divorced', 'Divorced'),
            ('widowed', 'Widowed'),
            ('separated', 'Separated'),
            ('prefer_not_to_say', 'Prefer not to say'),
        ],
        string='Marital Status',
    )

    selection_form_document_ids = fields.One2many(
        'hr.applicant.document', 'applicant_id',
        string='Selection Form Documents',
    )

    _PIPELINE_ORDER = (
        'applied', 'shortlisted', 'evaluation', 'submission',
        'contract', 'compliance', 'email_id', 'onboarded',
    )

    _STAGE_NAME_TO_PIPELINE = (
        (('onboarded', 'onboarding complete'), 'onboarded'),
        (('welcome mail', 'induction', 'email created', 'it email'), 'email_id'),
        (('statutory', 'compliance'), 'compliance'),
        (('contract', 'offer signed', 'hired'), 'contract'),
        (('selection form', 'form sent', 'form submit', 'form validat'), 'submission'),
        (('evaluation', 'interview'), 'evaluation'),
        (('shortlist',), 'shortlisted'),
        (('reject', 'refuse'), 'rejected'),
        (('applied', 'application', 'sourc', 'resume upload', 'screening'), 'applied'),
    )

    @classmethod
    def _derive_pipeline_status_from_name(cls, stage_name):
        if not stage_name:
            return None
        lower = stage_name.strip().lower()
        for keywords, pstatus in cls._STAGE_NAME_TO_PIPELINE:
            if any(k in lower for k in keywords):
                return pstatus
        return None

    @classmethod
    def _pipeline_order_index(cls, status):
        if not status or status == 'rejected':
            return -1
        try:
            return cls._PIPELINE_ORDER.index(status)
        except ValueError:
            return -1

    @api.model_create_multi
    def create(self, vals_list):
        for vals in vals_list:
            if not vals.get('candidate_code'):
                import uuid as _uuid
                yymm = fields.Datetime.now().strftime('%b%y').upper()
                vals['candidate_code'] = 'ETH-%s-%s' % (
                    yymm, _uuid.uuid4().hex[:6].upper(),
                )
        return super().create(vals_list)

    def write(self, vals):
        if 'job_id' in vals:
            History = self.env['hr.applicant.job.history'].sudo()
            new_job_id = vals.get('job_id') or False
            for applicant in self:
                old_job_id = applicant.job_id.id if applicant.job_id else False
                if old_job_id != new_job_id:
                    History.create({
                        'applicant_id': applicant.id,
                        'previous_job_id': old_job_id,
                        'new_job_id': new_job_id,
                        'changed_by_id': self.env.uid,
                    })
        result = super().write(vals)
        if vals.get('stage_id') and 'pipeline_status' not in vals:
            self._sync_pipeline_status_from_stage()
        if not self.env.context.get('skip_user_sync') and (
            'partner_name' in vals or 'partner_phone' in vals
        ):
            self._sync_user_from_applicant(vals)
        return result

    def _sync_user_from_applicant(self, vals):
        for applicant in self:
            if not applicant.candidate_user_id:
                continue
            sync = {}
            if 'partner_name' in vals:
                sync['name'] = vals['partner_name']
            if 'partner_phone' in vals:
                sync['phone'] = vals['partner_phone']
            if sync:
                applicant.candidate_user_id.sudo().with_context(
                    skip_applicant_sync=True,
                ).write(sync)

    def _sync_pipeline_status_from_stage(self):
        for applicant in self:
            if applicant.pipeline_status == 'rejected':
                continue
            derived = self._derive_pipeline_status_from_name(
                applicant.stage_id.name if applicant.stage_id else '',
            )
            if not derived:
                continue
            current = applicant.pipeline_status or 'applied'
            if derived == 'rejected':
                if current != 'rejected':
                    super(HrApplicant, applicant).write(
                        {'pipeline_status': 'rejected'},
                    )
                continue
            if self._pipeline_order_index(derived) > self._pipeline_order_index(current):
                super(HrApplicant, applicant).write(
                    {'pipeline_status': derived},
                )

    @api.onchange('candidate_id', 'active', 'stage_id')
    def _onchange_warn_active_application(self):
        if not self.candidate_id or not self.active:
            return
        if self.stage_id and self.stage_id.hired_stage:
            return
        if self.refuse_reason_id:
            return
        domain = [
            ('candidate_id', '=', self.candidate_id.id),
            ('active', '=', True),
            ('refuse_reason_id', '=', False),
        ]
        if isinstance(self.id, int):
            domain.append(('id', '!=', self.id))
        others = self.env['hr.applicant'].sudo().search(domain).filtered(
            lambda a: not (a.stage_id and a.stage_id.hired_stage)
        )
        if others:
            return {
                'warning': {
                    'title': "Existing Application",
                    'message': (
                        "Candidate '%s' already has an active application "
                        "(id %s) for job '%s'. Cancel it before applying to "
                        "another job."
                    ) % (
                        self.candidate_id.name,
                        others[0].id,
                        others[0].job_id.name or '?',
                    ),
                }
            }

    def action_screen_resume(self):
        self.ensure_one()
        from odoo.addons.ethara_hrms_extension.services import llm_client

        missing = []
        if not self.job_id:
            missing.append('Applied Job (job_id)')
        if not (self.resume_url or '').strip():
            missing.append('Resume URL (resume_url)')
        if missing:
            raise UserError(
                'Cannot screen resume — the following required field(s) are '
                'missing on this applicant:\n\n  • %s\n\n'
                'Please fill them in and try again.' % '\n  • '.join(missing)
            )

        ICP = self.env['ir.config_parameter'].sudo()
        api_key = (ICP.get_param('ethara_hrms.llm_api_key', '') or '').strip()
        if not api_key:
            raise UserError(
                'LLM API key not configured. Set the '
                '"ethara_hrms.llm_api_key" system parameter '
                '(Settings → Technical → System Parameters).'
            )
        base_url = ICP.get_param(
            'ethara_hrms.llm_base_url', llm_client.DEFAULT_BASE_URL,
        )
        model = ICP.get_param(
            'ethara_hrms.llm_model', llm_client.DEFAULT_LLM_MODEL,
        )
        timeout = int(ICP.get_param('ethara_hrms.llm_timeout', '60'))
        max_retries = int(ICP.get_param('ethara_hrms.llm_max_retries', '3'))

        resume_text = self._extract_resume_text()
        if not resume_text:
            raise UserError(
                'Could not extract text from the resume. Supported formats: '
                'PDF (with a text layer) and DOCX. If the file is a scanned '
                'image, an encrypted PDF, a legacy .doc, or another format '
                '(image, HTML, ZIP archive), please ask the candidate to '
                're-upload as a searchable PDF or DOCX. See the Odoo log '
                'for the exact detection result.'
            )

        wrong_kind = self._detect_non_resume(resume_text)
        if wrong_kind:
            raise UserError(
                'The uploaded file looks like a %s, not a resume. '
                'Please replace it with the candidate\'s resume (CV) '
                'in PDF or DOCX format before running screening.' % wrong_kind
            )

        system_prompt = (
            "You are an expert technical recruiter evaluating a candidate's "
            "resume for a job opening. Respond with a single JSON object "
            "containing exactly these keys: matchScore (integer 0-100), "
            "recommendation (one of: shortlist, reject, maybe, needs_review), "
            "summary (2-3 sentences), strengths (array of short strings), "
            "gaps (array of short strings). Do not include any prose "
            "outside the JSON object."
        )
        user_text = (
            "Job title: %s\n\nJob description:\n%s\n\n"
            "Candidate resume (first 6000 chars):\n%s"
        ) % (
            self._resume_job_title(),
            self._resume_job_description(),
            resume_text[:6000],
        )

        self.write({'resume_llm_status': 'processing', 'resume_llm_error': False})
        try:
            result = llm_client.chat_completion(
                api_key=api_key,
                model=model,
                system_prompt=system_prompt,
                user_text=user_text,
                temperature=0.2,
                base_url=base_url,
                timeout=timeout,
                max_retries=max_retries,
                app_title='Ethara HRMS Resume Screening',
            )
        except Exception as exc:
            _logger.exception(
                'Resume screening failed for hr.applicant %s', self.id,
            )
            self.write({
                'resume_llm_status': 'failed',
                'resume_llm_error': str(exc),
            })
            raise UserError('Screening failed: %s' % exc) from exc

        payload = self._parse_screening_json(result.get('content') or '')
        strengths = payload.get('strengths') or []
        gaps = payload.get('gaps') or []
        raw_score = (
            payload.get('matchScore')
            if payload.get('matchScore') is not None
            else payload.get('match_score')
            if payload.get('match_score') is not None
            else payload.get('score')
        )
        raw_summary = payload.get('summary')
        if isinstance(raw_summary, (dict, list)):
            raw_summary = json.dumps(raw_summary)
        self.write({
            'resume_score': self._clamp_score(raw_score),
            'resume_recommendation': self._normalize_recommendation(
                payload.get('recommendation'),
            ),
            'resume_summary': (raw_summary or '').strip(),
            'resume_strengths': '\n'.join(
                '- %s' % s for s in strengths if s
            ),
            'resume_gaps': '\n'.join('- %s' % s for s in gaps if s),
            'resume_screening_payload': json.dumps(payload, indent=2),
            'resume_llm_status': 'completed',
            'resume_llm_error': False,
            'resume_llm_model_used': result.get('model') or model,
            'resume_llm_prompt_tokens': int(result.get('prompt_tokens') or 0),
            'resume_llm_completion_tokens': int(
                result.get('completion_tokens') or 0,
            ),
            'resume_llm_latency_ms': int(result.get('latency_ms') or 0),
            'resume_screened_at': fields.Datetime.now(),
            'resume_text': resume_text,
        })
        self._advance_stage_after_screening()
        return {
            'type': 'ir.actions.client',
            'tag': 'display_notification',
            'params': {
                'title': 'Resume screened',
                'message': 'Verdict: %s · Score: %s/100' % (
                    self.resume_recommendation or '?',
                    int(self.resume_score or 0),
                ),
                'type': 'success',
                'sticky': False,
            },
        }

    def _extract_resume_text(self):
        self.ensure_one()
        url = (self.resume_url or '').strip()
        if not url:
            return ''
        return self._extract_resume_text_from_url(url)

    def _extract_resume_text_from_url(self, url):
        pdf_bytes = self._http_get_pdf(url)
        if not pdf_bytes:
            return ''
        try:
            return self._pdf_bytes_to_text(pdf_bytes)
        except Exception as exc:
            _logger.warning(
                'Resume PDF parse failed for %s: %s', url, exc,
            )
            return ''

    @staticmethod
    def _http_get_pdf(url):
        try:
            import requests
            resp = requests.get(
                url,
                headers={'User-Agent': 'Ethara-HRMS-Screening/1.0'},
                timeout=30,
                allow_redirects=True,
            )
            resp.raise_for_status()
            return resp.content
        except Exception as exc:
            _logger.info(
                'Resume URL fetch via requests failed (%s); trying curl fallback',
                exc,
            )
        import shutil
        import subprocess
        import tempfile
        curl = shutil.which('curl')
        if not curl:
            _logger.warning('curl not on PATH — cannot fetch %s', url)
            return b''
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tf:
            tmp_path = tf.name
        try:
            result = subprocess.run(
                [
                    curl, '-sSL', '--max-time', '30',
                    '-A', 'Ethara-HRMS-Screening/1.0',
                    '-o', tmp_path, url,
                ],
                capture_output=True, timeout=35, check=False,
            )
            if result.returncode != 0:
                _logger.warning(
                    'curl fetch failed for %s: rc=%s stderr=%s',
                    url, result.returncode,
                    (result.stderr or b'').decode(errors='replace')[:200],
                )
                return b''
            with open(tmp_path, 'rb') as fh:
                return fh.read()
        except Exception as exc:
            _logger.warning('curl fetch exception for %s: %s', url, exc)
            return b''
        finally:
            try:
                import os
                os.unlink(tmp_path)
            except OSError:
                pass

    _NON_RESUME_MARKERS = {
        'Aadhaar / UIDAI card': (
            'aadhaar', 'आधार', 'uidai',
            'unique identification authority',
        ),
        'PAN card': (
            'permanent account number', 'income tax department',
            'भारत सरकार आयकर विभाग',
        ),
        'Passport': (
            'passport no', 'republic of india passport',
            'ministry of external affairs',
        ),
        'Cancelled cheque or bank document': (
            'or bearer', 'a/c payee only', 'ifsc code',
            'account holder',
        ),
        'Academic marksheet or degree certificate': (
            'central board of secondary education',
            'senior secondary school certificate',
            'university of', 'has been awarded',
            'roll no.', 'roll number',
        ),
        'Driving licence': (
            'driving licence', 'transport department',
            'motor vehicles',
        ),
    }
    _RESUME_MARKERS = (
        'experience', 'skills', 'education', 'employment',
        'resume', 'curriculum vitae', 'objective', 'profile',
        'summary', 'responsibilities', 'worked as', 'developed',
        'implemented', 'managed', 'proficient', 'expertise',
        'certifications', 'projects', 'internship', 'linkedin',
        'github',
    )

    @classmethod
    def _detect_non_resume(cls, text):
        lower = (text or '').lower()
        resume_hits = sum(1 for m in cls._RESUME_MARKERS if m in lower)
        for kind, markers in cls._NON_RESUME_MARKERS.items():
            wrong_hits = sum(1 for m in markers if m in lower)
            if wrong_hits >= 2 and wrong_hits > resume_hits:
                _logger.info(
                    'Rejected non-resume upload: kind=%s wrong_hits=%d '
                    'resume_hits=%d', kind, wrong_hits, resume_hits,
                )
                return kind
        return None

    @staticmethod
    def _pdf_bytes_to_text(file_bytes):
        if not file_bytes:
            return ''
        magic = file_bytes[:8]
        if magic.startswith(b'%PDF'):
            try:
                try:
                    from pypdf import PdfReader
                except ImportError:
                    from PyPDF2 import PdfReader
                reader = PdfReader(BytesIO(file_bytes))
                text = '\n'.join(
                    (p.extract_text() or '') for p in reader.pages
                )
                if text.strip():
                    return text[:20000]
                _logger.info(
                    'PDF has no extractable text layer (scanned image?); '
                    'no OCR configured — returning empty',
                )
                return ''
            except Exception as exc:
                _logger.warning('PDF parse failed: %s', exc)
                return ''
        if magic.startswith(b'PK\x03\x04'):
            try:
                from docx import Document as _DocxDocument
                doc = _DocxDocument(BytesIO(file_bytes))
                paragraphs = [p.text for p in doc.paragraphs if p.text]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text:
                                paragraphs.append(cell.text)
                text = '\n'.join(paragraphs)
                return text[:20000] if text.strip() else ''
            except Exception as exc:
                _logger.warning('DOCX parse failed: %s', exc)
                return ''
        if magic.startswith(b'\xd0\xcf\x11\xe0'):
            _logger.warning(
                'Legacy .doc (OLE) format detected — not supported; '
                'ask the candidate to re-upload as PDF or DOCX',
            )
            return ''
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            return text[:20000] if text.strip() else ''
        except Exception:
            return ''

    def _resume_job_title(self):
        self.ensure_one()
        return (
            (self.job_id.name if self.job_id else None)
            or self.name
            or 'Unknown role'
        )

    def _resume_job_description(self):
        self.ensure_one()
        job = self.job_id
        pieces = []
        if job:
            if getattr(job, 'screening_prompt', None):
                pieces.append(job.screening_prompt)
            if job.description:
                pieces.append(
                    re.sub(
                        r'\s+', ' ',
                        unescape(re.sub(r'<[^>]+>', ' ', job.description)),
                    ).strip()
                )
        return '\n\n'.join(p for p in pieces if p) or (
            'Unspecified role — evaluate general fit.'
        )

    @staticmethod
    def _parse_screening_json(content):
        text = (content or '').strip()
        if text.startswith('```'):
            text = text.strip('`').strip()
            if text.lower().startswith('json'):
                text = text[4:].strip()
        try:
            return json.loads(text)
        except Exception:
            match = re.search(r'\{.*\}', text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(0))
                except Exception:
                    pass
        salvaged = {
            'summary': '',
            'recommendation': 'needs_review',
            'matchScore': 0,
            'strengths': [],
            'gaps': [],
        }
        m = re.search(r'"?matchScore"?\s*:\s*(\d+(?:\.\d+)?)', text)
        if m:
            try:
                salvaged['matchScore'] = int(float(m.group(1)))
            except (TypeError, ValueError):
                pass
        m = re.search(
            r'"?recommendation"?\s*:\s*"?(shortlist|reject|maybe|needs_review)"?',
            text,
            re.IGNORECASE,
        )
        if m:
            salvaged['recommendation'] = m.group(1).lower()
        m = re.search(
            r'"?summary"?\s*:\s*(.+?)(?=\s*"?(strengths|gaps|matchScore|recommendation)"?\s*:|\}|$)',
            text,
            re.DOTALL | re.IGNORECASE,
        )
        if m:
            summary_text = m.group(1).strip().rstrip(',').strip()
            if summary_text.startswith('"'):
                summary_text = summary_text[1:]
            if summary_text.endswith('"'):
                summary_text = summary_text[:-1]
            salvaged['summary'] = summary_text.strip()
        else:
            salvaged['summary'] = text[:500]
        return salvaged

    @staticmethod
    def _clamp_score(value):
        try:
            score = float(value or 0)
        except (TypeError, ValueError):
            score = 0.0
        return max(0.0, min(100.0, score))

    @staticmethod
    def _normalize_recommendation(value):
        v = (value or '').strip().lower()
        allowed = {'shortlist', 'reject', 'maybe', 'needs_review'}
        if v in allowed:
            return v
        if v in ('accept', 'yes', 'pass', 'proceed', 'hire', 'ship'):
            return 'shortlist'
        if v in ('no', 'fail', 'decline', 'block'):
            return 'reject'
        if v in ('hold', 'review', 'unsure'):
            return 'needs_review'
        return 'needs_review'

    def _advance_stage_after_screening(self, allow_regress=False):
        self.ensure_one()
        Stage = self.env['hr.recruitment.stage'].sudo()
        rec = self.resume_recommendation

        has_status = 'status' in self._fields

        if rec == 'reject':
            reject_vals = {}
            if self.pipeline_status != 'rejected':
                reject_vals['pipeline_status'] = 'rejected'
            if has_status:
                reject_vals['status'] = 'resume_screening_rejected'
            if reject_vals:
                self.sudo().write(reject_vals)
            return

        if rec != 'shortlist':
            return

        if has_status:
            self.sudo().write({'status': 'resume_screening_passed'})

        if not allow_regress and self.candidate_user_id:
            conflict = self.env['hr.applicant'].sudo().with_context(
                active_test=False,
            ).search(
                [
                    ('candidate_user_id', '=', self.candidate_user_id.id),
                    ('id', '!=', self.id),
                    ('active', '=', True),
                    ('refuse_reason_id', '=', False),
                    ('pipeline_status', '!=', 'rejected'),
                ],
                limit=1,
            )
            if conflict:
                _logger.info(
                    'LLM auto-shortlist skipped for applicant %s: '
                    'conflicting active application %s',
                    self.id, conflict.id,
                )
                return

        current_status = self.pipeline_status or 'applied'
        if allow_regress:
            if current_status != 'shortlisted':
                self.sudo().write({'pipeline_status': 'shortlisted'})
        else:
            if current_status == 'rejected':
                return
            if self._pipeline_order_index('shortlisted') > self._pipeline_order_index(current_status):
                self.sudo().write({'pipeline_status': 'shortlisted'})

        current_seq = self.stage_id.sequence if self.stage_id else -1
        job_ids = self.job_id.ids or [0]
        next_stage = Stage.search(
            [
                ('sequence', '>', current_seq),
                '|', ('job_ids', '=', False), ('job_ids', 'in', job_ids),
            ],
            order='sequence asc, id asc',
            limit=1,
        )
        if next_stage:
            self.sudo().write({'stage_id': next_stage.id})

    def _schedule_resume_screening(self):
        self.ensure_one()
        if not self.id:
            return
        dbname = self.env.cr.dbname
        applicant_id = self.id

        def _run_in_thread():
            try:
                registry = odoo.modules.registry.Registry(dbname)
                with registry.cursor() as cr:
                    thread_env = api.Environment(cr, SUPERUSER_ID, {})
                    applicant = thread_env['hr.applicant'].browse(applicant_id).exists()
                    if not applicant:
                        return
                    applicant.action_screen_resume()
            except Exception:
                _logger.exception(
                    'Async resume screening failed for applicant %s',
                    applicant_id,
                )

        def _fire_thread():
            threading.Thread(
                target=_run_in_thread,
                name='ethara-resume-screen-%s' % applicant_id,
                daemon=True,
            ).start()

        self.env.cr.postcommit.add(_fire_thread)
