"""Take-home / FSD assessment (P2-8): a side artifact, candidate.state UNTOUCHED.

Contract under test (models/iris_assessment.py + services/submission_extractor.py):

* creation guard: only candidates past screening (shipped / interview_ready /
  interviewed / scored) may receive an assessment;
* own ``status`` lifecycle (draft → sent → submitted → reviewed) independent
  of both ``candidate.state`` and ``llm_status``;
* submission extraction dispatch: PDF (magic-byte authoritative), docx
  (soft dependency — skipUnless), md/txt (UTF-8 errors="replace"),
  unsupported → clean ValueError/UserError;
* LOAD-BEARING: the LLM review draft changes NEITHER the assessment status
  NOR the candidate state — drafts never auto-apply;
* ``action_apply_draft`` fills EMPTY fields only and never raises on
  draft-format drift (partial-parse safe);
* ``action_finalize_review`` requires submitted + rating + recommendation +
  summary and renders the Feedback.md-shaped attachment in section order;
* full REST round-trip over ``/api/v1/iris/assessments``.
"""

import base64
import io
import json
import unittest
from datetime import timedelta

from odoo import fields
from odoo.exceptions import UserError, ValidationError
from odoo.tests.common import HttpCase, tagged

from .common import (
    API_KEY_PARAM,
    RESUME_TEXT,
    VALID_ASSESSMENT_DRAFT,
    VALID_SHIP_RECORD,
    IrisCase,
    make_pdf_bytes,
    mock_llm,
    patch_encryption_env,
)
from odoo.addons.api_auth_gateway.models import access_token as access_token_model
from odoo.addons.iris.models import credential_manager
from odoo.addons.iris.services import submission_extractor

try:
    from docx import Document as DocxDocument
    HAS_PYTHON_DOCX = True
except ImportError:
    DocxDocument = None
    HAS_PYTHON_DOCX = False

JSON_HEADERS = {"Content-Type": "application/json"}

SUBMISSION_MD = b"# Submission\n\nEval harness with retry logic and 14 tests."


@tagged("post_install", "-at_install", "iris")
class TestAssessmentFlow(IrisCase):
    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------
    def _shipped_candidate(self, name="Jane Doe"):
        candidate = self._make_candidate(name=name)
        self._screen(candidate, self.VALID_SHIP_RECORD)
        self.assertEqual(candidate.state, "shipped")
        return candidate

    def _upload_md(self, assessment, content, filename="submission.md"):
        assessment.write({
            "submission_file": base64.b64encode(content).decode("ascii"),
            "submission_filename": filename,
        })

    def _submitted_assessment(self, candidate=None):
        candidate = candidate or self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        assessment.action_send_brief()
        self._upload_md(assessment, SUBMISSION_MD)
        self.assertEqual(assessment.status, "submitted")
        return assessment

    def _drafted_assessment(self, candidate=None):
        assessment = self._submitted_assessment(candidate)
        assessment.action_generate_review_draft()
        with mock_llm(self.VALID_ASSESSMENT_DRAFT):
            self._run_llm_queue()
        self.assertEqual(assessment.llm_status, "done")
        return assessment

    def _reviewed_assessment(self):
        assessment = self._submitted_assessment()
        assessment.write({
            "rating": "above_average",
            "recommendation": "lean_hire",
            "summary": "Solid execution under the brief.",
        })
        assessment.action_finalize_review()
        self.assertEqual(assessment.status, "reviewed")
        return assessment

    # ------------------------------------------------------------------
    # Creation guard
    # ------------------------------------------------------------------
    def test_creation_guard_per_candidate_state(self):
        blocked = (
            "draft", "screening", "hold", "blocked", "pending_block",
            "needs_review", "hired", "rejected",
        )
        for state in blocked:
            with self.subTest(state=state):
                candidate = self._make_candidate(name=f"Guard {state}")
                candidate.write({"state": state})
                with self.assertRaises(ValidationError):
                    self._make_assessment(candidate)
                self.env.invalidate_all()

        allowed = ("shipped", "interview_ready", "interviewed", "scored")
        for state in allowed:
            with self.subTest(state=state):
                candidate = self._make_candidate(name=f"Allowed {state}")
                candidate.write({"state": state})
                assessment = self._make_assessment(candidate)
                self.assertEqual(assessment.status, "draft")
                self.assertEqual(assessment.llm_status, "none")

    # ------------------------------------------------------------------
    # Status transitions
    # ------------------------------------------------------------------
    def test_send_brief_transitions_and_guards(self):
        candidate = self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        self.assertEqual(assessment.status, "draft")

        assessment.write({"brief": False})
        with self.assertRaises(UserError):
            assessment.action_send_brief()
        self.env.invalidate_all()
        self.assertEqual(assessment.status, "draft")

        assessment.write({"brief": "Build a small eval harness."})
        assessment.action_send_brief()
        self.assertEqual(assessment.status, "sent")
        self.assertTrue(assessment.sent_at)

        # Only from draft.
        with self.assertRaises(UserError):
            assessment.action_send_brief()

    def test_submission_upload_moves_to_submitted(self):
        assessment = self._submitted_assessment()
        self.assertTrue(assessment.submitted_at)
        self.assertIn("retry logic", assessment.submission_text)

        # A re-upload while submitted re-extracts in place.
        self._upload_md(assessment, b"# Submission v2\n\nNow with metrics.")
        self.assertEqual(assessment.status, "submitted")
        self.assertIn("Now with metrics.", assessment.submission_text)

    def test_submission_upload_from_draft_also_submits(self):
        candidate = self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        self._upload_md(assessment, SUBMISSION_MD)
        self.assertEqual(assessment.status, "submitted")

    def test_submission_locked_after_review(self):
        assessment = self._reviewed_assessment()
        with self.assertRaises(UserError):
            self._upload_md(assessment, b"late upload")
        self.env.invalidate_all()
        self.assertEqual(assessment.status, "reviewed")

    # ------------------------------------------------------------------
    # Extractor dispatch (service level)
    # ------------------------------------------------------------------
    def test_extractor_dispatch_pdf(self):
        pdf = make_pdf_bytes("FSD submission — eval harness design")
        text = submission_extractor.extract_submission_text(
            pdf, "submission.pdf",
        )
        self.assertIn("eval harness design", text)

    def test_extractor_pdf_magic_wins_without_extension(self):
        pdf = make_pdf_bytes("magic-byte dispatch")
        text = submission_extractor.extract_submission_text(pdf, "")
        self.assertIn("magic-byte dispatch", text)

    def test_extractor_dispatch_md_and_txt(self):
        for filename in ("notes.md", "notes.txt"):
            with self.subTest(filename=filename):
                text = submission_extractor.extract_submission_text(
                    b"# Title\n\nBody text.", filename,
                )
                self.assertEqual(text, "# Title\n\nBody text.")

    def test_extractor_txt_replaces_undecodable_bytes(self):
        text = submission_extractor.extract_submission_text(
            b"ok \xff\xfe bytes", "raw.txt",
        )
        self.assertIn("ok", text)
        self.assertIn("�", text)

    @unittest.skipUnless(HAS_PYTHON_DOCX, "python-docx is not installed")
    def test_extractor_dispatch_docx(self):
        buffer = io.BytesIO()
        document = DocxDocument()
        document.add_paragraph("Take-home submission body.")
        document.add_paragraph("Second paragraph.")
        document.save(buffer)
        text = submission_extractor.extract_submission_text(
            buffer.getvalue(), "submission.docx",
        )
        self.assertIn("Take-home submission body.", text)
        self.assertIn("Second paragraph.", text)

    def test_extractor_docx_without_zip_magic_rejected(self):
        # Magic check fires before the soft import: runs without python-docx.
        with self.assertRaises(ValueError):
            submission_extractor.extract_submission_text(
                b"not a zip archive", "fake.docx",
            )

    def test_extractor_unsupported_and_empty_rejected(self):
        with self.assertRaises(ValueError):
            submission_extractor.extract_submission_text(
                b"PK\x03\x04zipdata", "archive.zip",
            )
        with self.assertRaises(ValueError):
            submission_extractor.extract_submission_text(b"", "empty.md")

    def test_unsupported_submission_surfaces_as_usererror(self):
        candidate = self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        with self.assertRaises(UserError):
            self._upload_md(assessment, b"random bytes", filename="notes.rtf")
        self.env.invalidate_all()
        self.assertEqual(assessment.status, "draft")
        self.assertFalse(assessment.submission_text)

    def test_pdf_submission_extracts_through_the_model(self):
        candidate = self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        assessment.write({
            "submission_file": base64.b64encode(
                make_pdf_bytes("PDF take-home: throughput design doc"),
            ).decode("ascii"),
            "submission_filename": "takehome.pdf",
        })
        self.assertEqual(assessment.status, "submitted")
        self.assertIn("throughput design doc", assessment.submission_text)

    # ------------------------------------------------------------------
    # LLM review draft — LOAD-BEARING: no state machine moves
    # ------------------------------------------------------------------
    def test_review_draft_guards(self):
        candidate = self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        # Draft status: nothing submitted yet.
        with self.assertRaises(UserError):
            assessment.action_generate_review_draft()
        self.env.invalidate_all()

        # Submitted status but no extracted text.
        assessment.write({"status": "submitted"})
        with self.assertRaises(UserError):
            assessment.action_generate_review_draft()
        self.env.invalidate_all()

        # Already in flight.
        submitted = self._submitted_assessment()
        submitted.write({"llm_status": "queued"})
        with self.assertRaises(UserError):
            submitted.action_generate_review_draft()

    def test_review_draft_changes_neither_status_nor_candidate_state(self):
        """LOAD-BEARING: the LLM output is a DRAFT — no state machine moves."""
        candidate = self._shipped_candidate()
        assessment = self._submitted_assessment(candidate)

        assessment.action_generate_review_draft()
        self.assertEqual(assessment.llm_status, "queued")
        self.assertEqual(assessment.status, "submitted")
        self.assertEqual(candidate.state, "shipped")

        with mock_llm(self.VALID_ASSESSMENT_DRAFT):
            self._run_llm_queue()

        self.assertEqual(assessment.llm_status, "done")
        self.assertEqual(
            assessment.llm_draft_markdown, self.VALID_ASSESSMENT_DRAFT,
        )
        # NEITHER the assessment status NOR the candidate state moved.
        self.assertEqual(assessment.status, "submitted")
        self.assertEqual(candidate.state, "shipped")
        # Structured fields untouched: drafts never auto-apply.
        self.assertFalse(assessment.rating)
        self.assertFalse(assessment.recommendation)
        self.assertFalse(assessment.summary)

        date_str = fields.Date.context_today(assessment).isoformat()
        self.assertEqual(
            assessment.draft_attachment_id.name,
            f"assessment-draft-doe-{date_str}.md",
        )

        # Prompt carries brief + submission + resume (seniority calibration).
        prompt = assessment.llm_prompt_input
        self.assertIn("ASSESSMENT BRIEF:", prompt)
        self.assertIn("CANDIDATE SUBMISSION", prompt)
        self.assertIn(
            "CANDIDATE RESUME (context for seniority calibration):", prompt,
        )
        self.assertIn(self.RESUME_TEXT, prompt)

    def test_review_draft_failure_leaves_everything_unchanged(self):
        candidate = self._shipped_candidate()
        assessment = self._submitted_assessment(candidate)
        assessment.action_generate_review_draft()
        with mock_llm(side_effect=Exception("draft boom")):
            self._run_llm_queue()

        self.assertEqual(assessment.llm_status, "failed")
        self.assertIn("draft boom", assessment.llm_error)
        self.assertEqual(assessment.status, "submitted")
        self.assertEqual(candidate.state, "shipped")
        self.assertFalse(assessment.llm_draft_markdown)

        # Retry is failed-only, then re-queues.
        assessment.action_retry_llm()
        self.assertEqual(assessment.llm_status, "queued")
        with mock_llm(self.VALID_ASSESSMENT_DRAFT):
            self._run_llm_queue()
        self.assertEqual(assessment.llm_status, "done")
        with self.assertRaises(UserError):
            assessment.action_retry_llm()

    # ------------------------------------------------------------------
    # action_apply_draft
    # ------------------------------------------------------------------
    def test_apply_draft_fills_all_empty_fields(self):
        assessment = self._drafted_assessment()
        assessment.action_apply_draft()

        self.assertEqual(assessment.rating, "above_average")
        self.assertEqual(assessment.recommendation, "lean_hire")
        self.assertIn("complete, working submission", assessment.summary)
        self.assertIn("Error handling", assessment.strengths)
        self.assertIn("Observability", assessment.concerns)
        self.assertIn("throughput problem", assessment.fit_for_current_need)
        self.assertIn("Pair the first month", assessment.recommendation_conditions)
        # Copying the draft is a helper — the status is still untouched.
        self.assertEqual(assessment.status, "submitted")

    def test_apply_draft_keeps_existing_values(self):
        assessment = self._drafted_assessment()
        assessment.write({
            "summary": "HUMAN SUMMARY — do not overwrite.",
            "rating": "poor",
        })
        assessment.action_apply_draft()

        # Human edits always win; only empty fields were filled.
        self.assertEqual(assessment.summary, "HUMAN SUMMARY — do not overwrite.")
        self.assertEqual(assessment.rating, "poor")
        self.assertEqual(assessment.recommendation, "lean_hire")
        self.assertIn("Error handling", assessment.strengths)

    def test_apply_draft_partial_parse_never_raises(self):
        assessment = self._submitted_assessment()
        assessment.write({
            "llm_draft_markdown": "free-form text, no headings, no anchors",
        })
        assessment.action_apply_draft()  # must not raise
        self.assertFalse(assessment.rating)
        self.assertFalse(assessment.recommendation)
        self.assertFalse(assessment.summary)

        # A draft with only one recognisable section fills just that field.
        assessment.write({
            "llm_draft_markdown": "## Summary\n\nOnly a summary section.\n",
        })
        assessment.action_apply_draft()
        self.assertIn("Only a summary section.", assessment.summary)
        self.assertFalse(assessment.rating)
        self.assertFalse(assessment.recommendation)

    def test_apply_draft_guards(self):
        assessment = self._submitted_assessment()
        with self.assertRaises(UserError):
            assessment.action_apply_draft()  # no draft yet
        self.env.invalidate_all()

        reviewed = self._reviewed_assessment()
        reviewed.write({"llm_draft_markdown": self.VALID_ASSESSMENT_DRAFT})
        with self.assertRaises(UserError):
            reviewed.action_apply_draft()  # locked after review

    # ------------------------------------------------------------------
    # action_finalize_review
    # ------------------------------------------------------------------
    def test_finalize_requires_rating_recommendation_summary(self):
        assessment = self._submitted_assessment()
        with self.assertRaises(UserError):
            assessment.action_finalize_review()
        self.env.invalidate_all()

        assessment.write({
            "rating": "above_average", "recommendation": "lean_hire",
        })
        with self.assertRaises(UserError):
            assessment.action_finalize_review()  # summary still missing
        self.env.invalidate_all()

        assessment.write({"summary": "Solid."})
        assessment.action_finalize_review()
        self.assertEqual(assessment.status, "reviewed")

    def test_finalize_requires_submitted_status(self):
        candidate = self._shipped_candidate()
        assessment = self._make_assessment(candidate)
        assessment.write({
            "rating": "average", "recommendation": "hire", "summary": "x",
        })
        with self.assertRaises(UserError):
            assessment.action_finalize_review()  # still draft

    def test_finalize_renders_feedback_attachment_in_order(self):
        candidate = self._shipped_candidate()
        assessment = self._drafted_assessment(candidate)
        assessment.action_apply_draft()
        assessment.action_finalize_review()

        self.assertEqual(assessment.status, "reviewed")
        self.assertEqual(assessment.reviewed_by, self.env.user)
        self.assertTrue(assessment.reviewed_at)
        # Finalizing the side artifact never advances the candidate.
        self.assertEqual(candidate.state, "shipped")

        attachment = assessment.attachment_id
        date_str = fields.Date.context_today(assessment).isoformat()
        self.assertEqual(
            attachment.name, f"assessment-feedback-doe-{date_str}.md",
        )
        content = base64.b64decode(attachment.datas).decode("utf-8")

        self.assertIn("# Assessment Feedback — Jane Doe", content)
        self.assertIn("- **Rating:** Above Average", content)
        self.assertIn("- **Recommendation:** **Lean Hire**", content)

        # Feedback.md section order is fixed.
        headings = (
            "## Summary",
            "## Strengths",
            "## Concerns",
            "## Fit for Current Need",
            "## Recommendation — Lean Hire, with conditions",
        )
        for heading in headings:
            self.assertIn(heading, content)
        positions = [content.index(heading) for heading in headings]
        self.assertEqual(positions, sorted(positions))

        bodies = self._chatter_bodies(candidate)
        self.assertTrue(
            any("Assessment reviewed" in body for body in bodies),
            f"no review note on the candidate chatter: {bodies}",
        )


@tagged("post_install", "-at_install", "iris")
class TestAssessmentApi(HttpCase):
    """REST round-trip for ``/api/v1/iris/assessments`` (HttpCase)."""

    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        env_patcher, _key = patch_encryption_env()
        cls.addClassCleanup(env_patcher.stop)

        Users = cls.env["res.users"].with_context(no_reset_password=True)
        base_group = cls.env.ref("base.group_user")
        cls.user_iris = Users.create({
            "name": "Assessment Api User",
            "login": "iris_assessment_api_user",
            "email": "iris_assessment_api_user@example.com",
            "group_ids": [(6, 0, [
                base_group.id,
                cls.env.ref("iris.group_iris_user").id,
            ])],
        })
        cls.user_manager = Users.create({
            "name": "Assessment Api Manager",
            "login": "iris_assessment_api_manager",
            "email": "iris_assessment_api_manager@example.com",
            "group_ids": [(6, 0, [
                base_group.id,
                cls.env.ref("iris.group_iris_manager").id,
            ])],
        })

        credential_manager.set_encrypted_param(
            cls.env, API_KEY_PARAM, "sk-iris-assessment-api-test",
        )
        connectors = cls.env["s3.connector"].sudo().search([])
        if connectors:
            connectors.unlink()
        cls.env.flush_all()

    def setUp(self):
        super().setUp()
        self.token_user = self._get_token("iris_assessment_api_user")
        self.token_manager = self._get_token("iris_assessment_api_manager")

    # ------------------------------------------------------------------
    # Helpers (same token/transport pattern as tests/test_api_controllers.py)
    # ------------------------------------------------------------------
    def _get_token(self, login):
        # api.access_token row created via ORM: the gateway's HTTP issuance
        # route commits + writes res_users_log, both forbidden inside the
        # test transaction. @validate_token only reads this table.
        user = self.env["res.users"].sudo().search(
            [("login", "=", login)], limit=1,
        )
        self.assertTrue(user, f"no such test user: {login}")
        token = access_token_model.nonce()
        self.env["api.access_token"].sudo().create({
            "user_id": user.id,
            "access_token": token,
            "refresh_token": access_token_model.nonce(),
            "expiry": fields.Datetime.now() + timedelta(seconds=3600),
        })
        self.env.flush_all()
        return token

    def _request(self, method, url, token=None, payload=None):
        headers = dict(JSON_HEADERS)
        if token:
            headers["access-token"] = token
        data = json.dumps(payload) if payload is not None else None
        self.env.flush_all()
        resp = self.url_open(url, data=data, headers=headers, method=method)
        self.env.invalidate_all()
        return resp

    def _make_candidate(self, name="Jane Doe", **overrides):
        vals = {"name": name, "resume_text": RESUME_TEXT}
        vals.update(overrides)
        candidate = self.env["iris.candidate"].create(vals)
        self.env.flush_all()
        return candidate

    def _run_llm_queue(self):
        self.env["iris.candidate"]._cron_process_llm_queue()
        self.env.flush_all()

    def _shipped_candidate(self, name="Jane Doe"):
        candidate = self._make_candidate(name=name)
        candidate.action_screen()
        with mock_llm(VALID_SHIP_RECORD):
            self._run_llm_queue()
        self.assertEqual(candidate.state, "shipped")
        return candidate

    # ------------------------------------------------------------------
    # Round-trip
    # ------------------------------------------------------------------
    def test_assessment_api_round_trip(self):
        candidate = self._shipped_candidate()
        cid = candidate.id

        # Creation guard surfaces as 400 for a draft candidate.
        draft = self._make_candidate(name="Still Draft")
        resp = self._request(
            "POST", f"/api/v1/iris/candidates/{draft.id}/assessments",
            token=self.token_user, payload={"brief": "x"},
        )
        self.assertEqual(resp.status_code, 400)

        # Create.
        resp = self._request(
            "POST", f"/api/v1/iris/candidates/{cid}/assessments",
            token=self.token_user,
            payload={"brief": "Build a small eval harness."},
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        body = resp.json()["assessment"]
        aid = body["id"]
        self.assertEqual(body["status"], "draft")
        self.assertEqual(body["candidate_id"], cid)

        # Unsupported submission type → clean 400.
        resp = self._request(
            "POST", f"/api/v1/iris/assessments/{aid}/submission",
            token=self.token_user,
            payload={
                "submission_base64": base64.b64encode(b"bytes").decode(),
                "submission_filename": "notes.rtf",
            },
        )
        self.assertEqual(resp.status_code, 400)

        # Markdown submission upload → submitted.
        resp = self._request(
            "POST", f"/api/v1/iris/assessments/{aid}/submission",
            token=self.token_user,
            payload={
                "submission_base64": base64.b64encode(SUBMISSION_MD).decode(),
                "submission_filename": "submission.md",
            },
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        self.assertEqual(resp.json()["assessment"]["status"], "submitted")

        # Review draft: queued over HTTP, completed by the queue cron.
        resp = self._request(
            "POST", f"/api/v1/iris/assessments/{aid}/review-draft",
            token=self.token_user,
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        self.assertEqual(resp.json()["llm_status"], "queued")
        with mock_llm(VALID_ASSESSMENT_DRAFT):
            self._run_llm_queue()

        resp = self._request(
            "GET", f"/api/v1/iris/assessments/{aid}", token=self.token_user,
        )
        body = resp.json()["assessment"]
        self.assertEqual(body["llm_status"], "done")
        self.assertIn("Assessment Review (DRAFT)", body["llm_draft_markdown"])
        # The draft moved neither state machine.
        self.assertEqual(body["status"], "submitted")
        candidate.invalidate_recordset()
        self.assertEqual(candidate.state, "shipped")

        # Invalid rating value → 400.
        resp = self._request(
            "PUT", f"/api/v1/iris/assessments/{aid}/feedback",
            token=self.token_user, payload={"rating": "stellar"},
        )
        self.assertEqual(resp.status_code, 400)

        # Structured feedback PUT.
        resp = self._request(
            "PUT", f"/api/v1/iris/assessments/{aid}/feedback",
            token=self.token_user,
            payload={
                "rating": "above_average",
                "recommendation": "lean_hire",
                "summary": "Solid execution under the brief.",
            },
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        body = resp.json()["assessment"]
        self.assertEqual(body["rating"], "above_average")
        self.assertEqual(body["recommendation"], "lean_hire")

        # Finalize → reviewed.
        resp = self._request(
            "POST", f"/api/v1/iris/assessments/{aid}/finalize",
            token=self.token_user,
        )
        self.assertEqual(resp.status_code, 200, resp.text)
        self.assertEqual(resp.json()["assessment"]["status"], "reviewed")

        # Feedback is locked after review.
        resp = self._request(
            "PUT", f"/api/v1/iris/assessments/{aid}/feedback",
            token=self.token_user, payload={"summary": "late edit"},
        )
        self.assertEqual(resp.status_code, 400)

        # List filters.
        resp = self._request(
            "GET",
            f"/api/v1/iris/assessments?candidate_id={cid}&status=reviewed",
            token=self.token_user,
        )
        self.assertEqual(resp.status_code, 200)
        self.assertEqual(
            [a["id"] for a in resp.json()["assessments"]], [aid],
        )

        # Manager-only delete.
        resp = self._request(
            "DELETE", f"/api/v1/iris/assessments/{aid}", token=self.token_user,
        )
        self.assertEqual(resp.status_code, 403)
        resp = self._request(
            "DELETE", f"/api/v1/iris/assessments/{aid}",
            token=self.token_manager,
        )
        self.assertEqual(resp.status_code, 200)
        resp = self._request(
            "GET", f"/api/v1/iris/assessments/{aid}", token=self.token_user,
        )
        self.assertEqual(resp.status_code, 404)

    def test_assessment_detail_404(self):
        resp = self._request(
            "GET", "/api/v1/iris/assessments/99999999", token=self.token_user,
        )
        self.assertEqual(resp.status_code, 404)
